"""Módulo para extraer texto de documentos subidos (TXT, PDF, DOCX)."""
import re
import io


def extract_text(uploaded_file) -> str:
    """Extrae texto de un archivo subido (UploadedFile de Streamlit)."""
    filename = uploaded_file.name.lower()
    raw_bytes = uploaded_file.read()
    uploaded_file.seek(0)

    if filename.endswith(".txt"):
        return _extract_txt(raw_bytes)
    elif filename.endswith(".pdf"):
        return _extract_pdf(raw_bytes)
    elif filename.endswith(".docx"):
        return _extract_docx(raw_bytes)
    else:
        raise ValueError(f"Formato no soportado: {filename}. Use TXT, PDF o DOCX.")


def extract_text_from_path(file_path: str) -> str:
    """Extrae texto de un archivo por ruta (para testing)."""
    path_lower = file_path.lower()
    with open(file_path, "rb") as f:
        raw_bytes = f.read()

    if path_lower.endswith(".txt"):
        return _extract_txt(raw_bytes)
    elif path_lower.endswith(".pdf"):
        return _extract_pdf(raw_bytes)
    elif path_lower.endswith(".docx"):
        return _extract_docx(raw_bytes)
    else:
        raise ValueError(f"Formato no soportado: {file_path}")


def _extract_txt(raw_bytes: bytes) -> str:
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            text = raw_bytes.decode(encoding)
            return _clean_text(text)
        except (UnicodeDecodeError, ValueError):
            continue
    return raw_bytes.decode("utf-8", errors="replace")


def _extract_pdf(raw_bytes: bytes) -> str:
    import pdfplumber
    pages_text = []
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
    return _clean_text("\n".join(pages_text))


def _extract_docx(raw_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(raw_bytes))
    full_text = []

    # 1. Extraer texto de párrafos
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # 2. Extraer texto de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)

    return _clean_text("\n".join(full_text))


def _clean_text(text: str) -> str:
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = text.strip()
    return text


def generate_summary(text: str, max_sentences: int = 8) -> str:
    """Genera un resumen básico del documento (primeras oraciones significativas)."""
    sentences = re.split(r'(?<=[.!?])\s+', text)
    meaningful = [s for s in sentences if len(s.split()) > 5]
    summary_sentences = meaningful[:max_sentences]
    if not summary_sentences:
        words = text.split()
        return " ".join(words[:200]) + ("..." if len(words) > 200 else "")
    return " ".join(summary_sentences)
