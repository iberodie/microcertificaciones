"""Módulo para generar el documento Word (.docx) con las recomendaciones."""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from templates.framework_text import (
    FRAMEWORK_INTRO, FRAMEWORK_SECTIONS, FRAMEWORK_SUMMARY
)


def generate_report(
    document_summary: str,
    competencies: list[dict],
    competencies_text: str,
    coursera_courses: list[dict],
    coursera_specializations: list[dict],
    external_results: list[dict],
    output_path: str,
    teacher_name: str = "Docente",
) -> str:
    """Genera el documento Word completo con todas las recomendaciones."""

    doc = Document()
    _setup_styles(doc)

    # === PORTADA ===
    _add_cover_page(doc, teacher_name)

    # === SECCIÓN 1: Resumen del Documento Base ===
    doc.add_page_break()
    _add_section_heading(doc, "1. Resumen del Documento Base")
    p = doc.add_paragraph(document_summary)
    p.style = doc.styles['Normal']

    # === SECCIÓN 2: Definición de Función/Habilidad(es) ===
    _add_section_heading(doc, "2. Definición de Función y Habilidades Identificadas")

    # Competencias extraídas
    sub = doc.add_paragraph()
    sub_run = sub.add_run("2.1 Competencias identificadas en el documento")
    sub_run.bold = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x8B, 0x0A, 0x1A)

    doc.add_paragraph(competencies_text)

    # Marco de microcredenciales
    sub2 = doc.add_paragraph()
    sub2_run = sub2.add_run("2.2 Marco conceptual de microcredenciales")
    sub2_run.bold = True
    sub2_run.font.size = Pt(13)
    sub2_run.font.color.rgb = RGBColor(0x8B, 0x0A, 0x1A)

    doc.add_paragraph(FRAMEWORK_INTRO)
    for section_data in FRAMEWORK_SECTIONS.values():
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(section_data["titulo"])
        run_title.bold = True
        run_title.font.size = Pt(11)
        doc.add_paragraph(section_data["descripcion"])

    doc.add_paragraph(FRAMEWORK_SUMMARY).italic = True

    # === SECCIÓN 3: Microcredenciales de Coursera ===
    doc.add_page_break()
    _add_section_heading(doc, "3. Microcredenciales de Coursera Disponibles")

    if coursera_courses:
        doc.add_paragraph(
            f"Se identificaron {len(coursera_courses)} cursos individuales y "
            f"{len(coursera_specializations)} especializaciones/certificados relevantes "
            f"en el catálogo de Coursera Enterprise."
        )

        # Cursos individuales
        if coursera_courses:
            sub3 = doc.add_paragraph()
            sub3_run = sub3.add_run("3.1 Cursos individuales recomendados")
            sub3_run.bold = True
            sub3_run.font.size = Pt(13)
            sub3_run.font.color.rgb = RGBColor(0x8B, 0x0A, 0x1A)

            for i, course in enumerate(coursera_courses, 1):
                _add_coursera_course(doc, course, i)

        # Especializaciones
        if coursera_specializations:
            sub4 = doc.add_paragraph()
            sub4_run = sub4.add_run("3.2 Especializaciones y Certificados Profesionales")
            sub4_run.bold = True
            sub4_run.font.size = Pt(13)
            sub4_run.font.color.rgb = RGBColor(0x8B, 0x0A, 0x1A)

            for i, spec in enumerate(coursera_specializations, 1):
                _add_coursera_specialization(doc, spec, i)
    else:
        _add_no_results_message(doc, "Coursera")

    # === SECCIÓN 4: Microcertificaciones Externas ===
    doc.add_page_break()
    _add_section_heading(doc, "4. Microcertificaciones Externas (Fuera de Coursera)")

    if external_results:
        industry = [r for r in external_results if r.get("tipo") == "Industria"]
        platforms = [r for r in external_results if r.get("tipo") == "Plataforma"]

        if industry:
            sub5 = doc.add_paragraph()
            sub5_run = sub5.add_run("4.1 Certificaciones de Industria")
            sub5_run.bold = True
            sub5_run.font.size = Pt(13)
            sub5_run.font.color.rgb = RGBColor(0x8B, 0x0A, 0x1A)

            for i, cert in enumerate(industry, 1):
                _add_external_certification(doc, cert, i)

        if platforms:
            sub6 = doc.add_paragraph()
            sub6_run = sub6.add_run("4.2 Plataformas Educativas Recomendadas")
            sub6_run.bold = True
            sub6_run.font.size = Pt(13)
            sub6_run.font.color.rgb = RGBColor(0x8B, 0x0A, 0x1A)

            for i, plat in enumerate(platforms, 1):
                _add_external_certification(doc, plat, i)

    else:
        _add_no_results_message(doc, "externas")

    # === SECCIÓN 5: Propuestas Internas (si aplica) ===
    if not coursera_courses and not external_results:
        doc.add_page_break()
        _add_section_heading(doc, "5. Propuestas de Microcredenciales Internas")
        _add_internal_proposals(doc, competencies)

    # === PIE DE PÁGINA ===
    _add_footer(doc)

    # Guardar
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


# === Funciones auxiliares ===

def _setup_styles(doc):
    """Configura estilos del documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def _add_cover_page(doc, teacher_name: str):
    """Agrega la portada del documento."""
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("Recomendaciones de Microcredenciales")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x8B, 0x0A, 0x1A)  # Rojo Ibero

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = subtitle.add_run("Universidad Iberoamericana — Ciudad de México")
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run3 = info.add_run(f"Documento generado para: {teacher_name}")
    run3.font.size = Pt(12)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run4 = date_p.add_run(f"Fecha: {datetime.now().strftime('%d de %B de %Y')}")
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")
    run5 = disclaimer.add_run(
        "Este documento fue generado automáticamente por el Sistema de "
        "Recomendación de Microcredenciales de la Universidad Iberoamericana."
    )
    run5.font.size = Pt(9)
    run5.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run5.italic = True


def _add_section_heading(doc, text: str):
    """Agrega un encabezado de sección principal."""
    heading = doc.add_heading(text, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x8B, 0x0A, 0x1A)
        run.font.size = Pt(18)

    # Línea separadora
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


def _add_coursera_course(doc, course: dict, index: int):
    """Agrega un curso de Coursera al documento."""
    # Nombre del curso
    p_name = doc.add_paragraph()
    run_name = p_name.add_run(f"{index}. {course['nombre']}")
    run_name.bold = True
    run_name.font.size = Pt(12)
    run_name.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)

    # Tabla de info
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light List Accent 1'

    fields = [
        ("Institución/Partner", course.get("partner", "")),
        ("Duración", f"{course.get('horas', 'N/A')} horas" if course.get('horas') else "N/A"),
        ("Nivel", course.get("nivel", "")),
        ("Rating", f"{course.get('rating', 'N/A')}/5" if course.get('rating') else "N/A"),
        ("Dominio", f"{course.get('dominio', '')} — {course.get('subdominio', '')}"),
        ("Enlace", course.get("url", "")),
        ("Habilidades", course.get("skills", "")[:200]),
    ]

    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        # Formato label
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        for paragraph in table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

    # Justificación
    p_just = doc.add_paragraph()
    run_just_label = p_just.add_run("Justificación: ")
    run_just_label.bold = True
    run_just_label.font.size = Pt(10)
    run_just = p_just.add_run(course.get("justificacion", ""))
    run_just.font.size = Pt(10)
    run_just.italic = True

    doc.add_paragraph("")  # Espacio


def _add_coursera_specialization(doc, spec: dict, index: int):
    """Agrega una especialización de Coursera al documento."""
    p_name = doc.add_paragraph()
    run_name = p_name.add_run(f"{index}. {spec['nombre']}")
    run_name.bold = True
    run_name.font.size = Pt(12)
    run_name.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light List Accent 1'

    fields = [
        ("Institución/Partner", spec.get("partner", "")),
        ("Tipo", spec.get("tipo", "")),
        ("Número de cursos", str(spec.get("num_cursos", ""))),
        ("Nivel", spec.get("nivel", "")),
        ("Enlace", spec.get("url", "")),
    ]

    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        for paragraph in table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

    p_just = doc.add_paragraph()
    run_just_label = p_just.add_run("Justificación: ")
    run_just_label.bold = True
    run_just_label.font.size = Pt(10)
    run_just = p_just.add_run(spec.get("justificacion", ""))
    run_just.font.size = Pt(10)
    run_just.italic = True

    doc.add_paragraph("")


def _add_external_certification(doc, cert: dict, index: int):
    """Agrega una certificación externa al documento."""
    p_name = doc.add_paragraph()
    run_name = p_name.add_run(f"{index}. {cert['nombre']}")
    run_name.bold = True
    run_name.font.size = Pt(12)
    run_name.font.color.rgb = RGBColor(0x1A, 0x73, 0x38)  # Verde

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light List Accent 1'

    fields = [
        ("Plataforma", cert.get("plataforma", "")),
        ("Enlace de acceso", cert.get("url", "")),
        ("Duración", cert.get("duracion", "Variable")),
        ("Costo", cert.get("costo", "Consultar sitio web")),
        ("Características", cert.get("caracteristicas", "")),
        ("Descripción", cert.get("descripcion", "")[:250]),
    ]

    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        for paragraph in table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

    p_just = doc.add_paragraph()
    run_just_label = p_just.add_run("Justificación: ")
    run_just_label.bold = True
    run_just_label.font.size = Pt(10)
    run_just = p_just.add_run(cert.get("justificacion", ""))
    run_just.font.size = Pt(10)
    run_just.italic = True

    doc.add_paragraph("")


def _add_no_results_message(doc, source: str):
    """Agrega mensaje cuando no hay resultados."""
    p = doc.add_paragraph()
    run = p.add_run(
        f"No se encontraron microcredenciales disponibles en {source} que coincidan "
        f"directamente con las competencias identificadas en el documento del docente. "
        f"Esto puede deberse a que el área temática es muy específica o que las opciones "
        f"disponibles actualmente no cubren exactamente este perfil."
    )
    run.font.color.rgb = RGBColor(0x99, 0x55, 0x00)
    run.italic = True


def _add_internal_proposals(doc, competencies: list[dict]):
    """Agrega propuestas de microcredenciales internas cuando no hay opciones externas."""
    doc.add_paragraph(
        "Dado que no se encontraron opciones externas que cubran adecuadamente "
        "las competencias identificadas, se sugiere que la Universidad Iberoamericana "
        "diseñe microcredenciales internas siguiendo estos lineamientos:"
    )

    doc.add_paragraph(FRAMEWORK_INTRO)

    if competencies:
        p = doc.add_paragraph()
        run = p.add_run("Competencias sugeridas para microcredenciales internas:")
        run.bold = True

        for comp in competencies[:8]:
            term = comp["term"].title()
            doc.add_paragraph(
                f"• Microcredencial en \"{term}\": Diseñar un programa de 20-40 horas "
                f"que certifique resultados de aprendizaje específicos en esta área, "
                f"incluyendo evaluación práctica y evidencia de competencia.",
                style='List Bullet'
            )

    doc.add_paragraph(
        "Cada microcredencial interna debe incluir: resultados de aprendizaje "
        "claramente definidos, evaluación rigurosa, carga de trabajo especificada, "
        "y alineación con marcos de cualificación reconocidos."
    )


def _add_footer(doc):
    """Agrega pie de página al documento."""
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(
        "─" * 40 + "\n"
        "Documento generado por el Sistema de Recomendación de Microcredenciales\n"
        "Universidad Iberoamericana — Ciudad de México\n"
        f"Generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
